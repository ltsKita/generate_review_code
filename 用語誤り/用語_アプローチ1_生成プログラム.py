import docx
import re

# dataディレクトリのwordファイルを読み込み
doc = docx.Document('用語誤り/data/【用語誤りデータ(55個)】第12条_まとめ資料_本文.docx')

# 校閲結果を保存するwordファイル
result_doc = docx.Document()

# ①「とき」と「時」の校閲
for para in doc.paragraphs:
    for run in para.runs:
        text = run.text
        words = text.split()
        for word in words:
            if word == "とき":
                if re.search(r'\b時点\b', text):
                    word = "時"
                elif re.search(r'\b条件\b', text):
                    word = "とき"
            elif word == "時":
                if re.search(r'\b条件\b', text):
                    word = "とき"
                elif re.search(r'\b時点\b', text):
                    word = "時"
        result_doc.add_paragraph(' '.join([word +'' for word in words]))

# ②「ほか」を意味する語の校閲
for para in doc.paragraphs:
    for run in para.runs:
        text = run.text
        words = text.split()
        for word in words:
            if word == "他":
                if re.search(r'\b複合語\b', text):
                    continue
                result_doc.add_paragraph('ほか ')
            elif word == "外":
                if re.search(r'\b複合語\b', text):
                    continue
                result_doc.add_paragraph('ほか ')

# 校閲結果をwordファイルとして出力
result_doc.save('result.docx')